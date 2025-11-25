"""CV tailoring engine using AI."""

import json
from pathlib import Path
from typing import Optional

import anthropic
import structlog

from config.prompts import PROMPTS
from config.settings import settings, DATA_DIR
from src.models import Job

logger = structlog.get_logger()


class CVTailor:
    """Generate tailored CVs using AI."""

    def __init__(
        self,
        api_key: Optional[str] = None,
        model: Optional[str] = None,
        master_cv_path: Optional[Path] = None,
    ):
        """
        Initialize CV tailor.

        Args:
            api_key: Anthropic API key
            model: Claude model to use
            master_cv_path: Path to master CV JSON
        """
        self.api_key = api_key or settings.anthropic_api_key
        self.model = model or settings.claude_model
        self.client = anthropic.Anthropic(api_key=self.api_key)

        # Load master CV
        self.master_cv_path = master_cv_path or DATA_DIR / "master_cv.json"
        self.master_cv = self._load_master_cv()

    def _load_master_cv(self) -> dict:
        """Load master CV from JSON file."""
        if self.master_cv_path.exists():
            with open(self.master_cv_path) as f:
                return json.load(f)
        else:
            logger.warning("Master CV not found")
            return {}

    async def tailor_cv(self, job: Job) -> dict:
        """
        Generate a tailored CV for a job.

        Args:
            job: Job to tailor CV for

        Returns:
            Dict containing tailored CV content
        """
        logger.info(
            "Tailoring CV",
            title=job.title,
            company=job.company,
        )

        prompt = PROMPTS["cv_tailoring"].format(
            master_cv_json=json.dumps(self.master_cv, indent=2),
            job_title=job.title,
            company=job.company,
            job_description=job.description[:6000],
            key_requirements=", ".join(job.key_requirements) if job.key_requirements else "Not specified",
        )

        try:
            response = self.client.messages.create(
                model=self.model,
                max_tokens=3000,
                messages=[
                    {"role": "user", "content": prompt}
                ],
            )

            response_text = response.content[0].text
            tailored_cv = self._parse_response(response_text)

            logger.info("CV tailored successfully")
            return tailored_cv

        except Exception as e:
            logger.error(f"CV tailoring failed: {e}")
            return {}

    def _parse_response(self, response_text: str) -> dict:
        """Parse JSON response from Claude."""
        import re

        # Try direct JSON parse
        try:
            return json.loads(response_text)
        except json.JSONDecodeError:
            pass

        # Try extracting from code block
        json_match = re.search(r"```(?:json)?\s*([\s\S]*?)```", response_text)
        if json_match:
            try:
                return json.loads(json_match.group(1))
            except json.JSONDecodeError:
                pass

        # Try finding JSON object
        json_match = re.search(r"\{[\s\S]*\}", response_text)
        if json_match:
            try:
                return json.loads(json_match.group(0))
            except json.JSONDecodeError:
                pass

        return {}

    def generate_docx(self, tailored_cv: dict, job: Job) -> bytes:
        """
        Generate a DOCX file from tailored CV data.

        Args:
            tailored_cv: Tailored CV content
            job: Job for filename

        Returns:
            DOCX file as bytes
        """
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO

        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Header - Name
        personal = self.master_cv.get("personal", {})
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(personal.get("name", "Jevan Cousins"))
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact info
        contact_parts = [
            personal.get("email", ""),
            personal.get("phone", ""),
            personal.get("linkedin", ""),
        ]
        contact_para = doc.add_paragraph(" | ".join(filter(None, contact_parts)))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Profile
        if tailored_cv.get("profile"):
            doc.add_heading("Profile", level=1)
            doc.add_paragraph(tailored_cv["profile"])

        # Experience
        if tailored_cv.get("experience"):
            doc.add_heading("Experience", level=1)
            for exp in tailored_cv["experience"]:
                # Company and title
                exp_para = doc.add_paragraph()
                company_run = exp_para.add_run(f"{exp.get('company', '')} - {exp.get('title', '')}")
                company_run.bold = True

                # Location and dates
                loc_para = doc.add_paragraph()
                loc_para.add_run(f"{exp.get('location', '')} | {exp.get('dates', '')}")

                # Bullets
                for bullet in exp.get("bullets", []):
                    bullet_para = doc.add_paragraph(bullet, style="List Bullet")

        # Skills
        if tailored_cv.get("skills_to_highlight"):
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(", ".join(tailored_cv["skills_to_highlight"]))

        # Education (from master CV)
        if self.master_cv.get("education"):
            doc.add_heading("Education", level=1)
            for edu in self.master_cv["education"]:
                edu_para = doc.add_paragraph()
                edu_run = edu_para.add_run(f"{edu.get('institution', '')} - {edu.get('degree', '')}")
                edu_run.bold = True
                doc.add_paragraph(f"{edu.get('location', '')} | {edu.get('start_date', '')} - {edu.get('end_date', '')}")

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
