"""Cover letter generation using AI."""

import json
from pathlib import Path
from typing import Optional

import anthropic
import structlog

from config.prompts import PROMPTS
from config.settings import settings, DATA_DIR
from src.models import Job

logger = structlog.get_logger()


class CoverLetterGenerator:
    """Generate personalised cover letters using AI."""

    def __init__(
        self,
        api_key: Optional[str] = None,
        model: Optional[str] = None,
        master_cv_path: Optional[Path] = None,
    ):
        """
        Initialize cover letter generator.

        Args:
            api_key: Anthropic API key
            model: Claude model to use
            master_cv_path: Path to master CV JSON
        """
        self.api_key = api_key or settings.anthropic_api_key
        self.model = model or settings.claude_model
        self.client = anthropic.Anthropic(api_key=self.api_key)

        # Load master CV for candidate summary
        self.master_cv_path = master_cv_path or DATA_DIR / "master_cv.json"
        self.master_cv = self._load_master_cv()
        self.candidate_summary = self._create_candidate_summary()

    def _load_master_cv(self) -> dict:
        """Load master CV from JSON file."""
        if self.master_cv_path.exists():
            with open(self.master_cv_path) as f:
                return json.load(f)
        return {}

    def _create_candidate_summary(self) -> str:
        """Create candidate summary for prompts."""
        if not self.master_cv:
            return "Solution Architect with 3+ years experience in finance, expertise in Python, SQL, and data analytics."

        summary_parts = []

        if self.master_cv.get("personal"):
            summary_parts.append(f"Name: {self.master_cv['personal'].get('name', '')}")

        if self.master_cv.get("profiles"):
            default_profile = self.master_cv["profiles"].get("default", "")
            if default_profile:
                summary_parts.append(f"\nProfile: {default_profile}")

        if self.master_cv.get("experience"):
            summary_parts.append("\nKey Experience:")
            for exp in self.master_cv["experience"][:2]:
                summary_parts.append(f"- {exp.get('title')} at {exp.get('company')}")
                if exp.get("bullets"):
                    for bullet in exp["bullets"][:2]:
                        text = bullet.get("text", bullet) if isinstance(bullet, dict) else bullet
                        summary_parts.append(f"  • {text[:150]}...")

        return "\n".join(summary_parts)

    async def generate(self, job: Job, company_research: str = None) -> str:
        """
        Generate a cover letter for a job.

        Args:
            job: Job to write cover letter for
            company_research: Optional company research to include

        Returns:
            Cover letter text
        """
        logger.info(
            "Generating cover letter",
            title=job.title,
            company=job.company,
        )

        # Get company research if not provided
        if not company_research:
            company_research = await self._research_company(job.company)

        prompt = PROMPTS["cover_letter"].format(
            candidate_summary=self.candidate_summary,
            company=job.company,
            job_title=job.title,
            job_description=job.description[:4000],
            company_research=company_research,
        )

        try:
            response = self.client.messages.create(
                model=self.model,
                max_tokens=1500,
                messages=[
                    {"role": "user", "content": prompt}
                ],
            )

            cover_letter = response.content[0].text

            logger.info("Cover letter generated successfully")
            return cover_letter

        except Exception as e:
            logger.error(f"Cover letter generation failed: {e}")
            return ""

    async def _research_company(self, company: str) -> str:
        """
        Do basic company research for cover letter.

        Args:
            company: Company name

        Returns:
            Brief company research summary
        """
        prompt = f"""Provide a brief (2-3 sentences) summary of {company} that would be useful for a job application cover letter. Focus on:
- What the company does
- Any notable recent achievements or initiatives
- Company culture or values

Keep it factual and professional."""

        try:
            response = self.client.messages.create(
                model=self.model,
                max_tokens=500,
                messages=[
                    {"role": "user", "content": prompt}
                ],
            )
            return response.content[0].text
        except Exception:
            return f"{company} is a company in the financial services industry."

    def generate_docx(self, cover_letter: str, job: Job) -> bytes:
        """
        Generate a DOCX file from cover letter text.

        Args:
            cover_letter: Cover letter text
            job: Job for filename

        Returns:
            DOCX file as bytes
        """
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        from datetime import datetime

        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Header - contact info
        personal = self.master_cv.get("personal", {})

        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(personal.get("name", "Jevan Cousins"))
        name_run.bold = True
        name_run.font.size = Pt(14)

        # Contact details
        doc.add_paragraph(personal.get("email", ""))
        doc.add_paragraph(personal.get("phone", ""))
        doc.add_paragraph(personal.get("linkedin", ""))

        # Date
        doc.add_paragraph()
        doc.add_paragraph(datetime.now().strftime("%d %B %Y"))

        # Recipient (generic)
        doc.add_paragraph()
        doc.add_paragraph("Hiring Manager")
        doc.add_paragraph(job.company)

        # Subject
        doc.add_paragraph()
        subject_para = doc.add_paragraph()
        subject_run = subject_para.add_run(f"Re: Application for {job.title}")
        subject_run.bold = True

        # Body
        doc.add_paragraph()
        for paragraph in cover_letter.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
